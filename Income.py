from docx import Document

from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pymysql
from tkinter import *
#root=Tk()
import os
class Income:
    def __init__(self, master,uidai):
        master.geometry('400x400+600+200')
        master.title("Login Page")
        master.wm_iconbitmap('C:\\Users\\jerin\\OneDrive\\Desktop\\favicon.ico')
        self.uidai=uidai
        self.f = Frame(master, width=400, height=500,background='purple')
        self.f.place(x='0', y='0')
        self.l1 = Label(self.f, text=" Name: ", borderwidth=2, relief="solid", width=10,font='Courier 11')
        self.l1.place(x='10', y='30')
        self.t1 = Entry(self.f, width=20, relief="ridge",font='Courier 12')
        self.t1.place(x='170', y='30')

        self.l2 = Label(self.f, text=" Occupation:", borderwidth=2, relief="solid", width=15,font='Courier 11',)
        self.l2.place(x='10', y='60')
        self.t2 = Entry(self.f, width=20, relief="ridge",font='Courier 12')
        self.t2.place(x='170', y='60')

        self.l3 = Label(self.f, text="Salary :", borderwidth=2, relief="solid", width=15,font='Courier 11')
        self.l3.place(x='10', y='90')
        self.t3 = Entry(self.f, width=20, relief="ridge",font='Courier 12')
        self.t3.place(x='170', y='90')

        self.l4 = Label(self.f, text=" Type starting and ending year using '-' !", borderwidth=2, relief="solid", width=40,font='Courier 11')
        self.l4.place(x='10', y='120')
        #self.t4 = Entry(self.f, width=20, relief="ridge",font='Courier 12')
        #self.t4.place(x='170', y='120')

        self.l5 = Label(self.f, text=" Year:", borderwidth=2, relief="solid", width=10,font='Courier 11')
        self.l5.place(x='10', y='150')
        self.t5 = Entry(self.f, width=20, relief="ridge",font='Courier 12')
        self.t5.place(x='170', y='150')

        self.b1=Button(self.f,text="SUBMIT",relief="raised",width=10)
        self.b1.bind("<Button-1>", self.income)
        self.b1.place(x='160', y='270')

    def income(self,event):

        name=self.t1.get()
        occup=self.t2.get()
        sal=self.t3.get()
        vyear=self.t5.get()



        document = Document()
        document.add_picture('project.png', width=Inches(0.8))

        document.add_paragraph('GOVERNMENT OF MAHARASHTRA')

        h = document.add_heading('INCOME CERTIFICATE', 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        i = document.add_paragraph('THIS   IS   TO  CERTIFY  THAT ')
        i.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = document.add_paragraph()
        p.add_run({name}).bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        q = document.add_paragraph('IS OCCUPIED AS  ')
        q.add_run({occup}).bold = True
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER

        r = document.add_paragraph('AND HAS AN ANNUAL INCOME OF   ')
        r.add_run({sal}).bold = True
        r.alignment = WD_ALIGN_PARAGRAPH.CENTER

        s = document.add_paragraph('IN   THE    YEAR   OF	')
        s.add_run({vyear}).bold = True
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_paragraph()

        document.add_picture('stamp.jpg', width=Inches(0.8))
        document.add_picture('sign.jpg', width=Inches(0.8))
        document.add_paragraph('NAYAB TAHSILDAR')
        document.add_paragraph('MULUND')

        sj = self.uidai + "i.docx"
        document.save(sj)
        #SQL ACTIVITY

        db = pymysql.connect("localhost", "root", "root", "adv")
        cursor = db.cursor()
        sql = "UPDATE cia SET INCOME = 1 WHERE UIDAI={} ".format(self.uidai)
        try:
            cursor.execute(sql)
            db.commit()
            os.startfile(str(self.uidai) + "i.docx")
            os.startfile(str(self.uidai) + "i.docx", "print")
            exit(0)
        except pymysql.Error as e:
            print(e)



#b = Income(root,'12345')
#root.mainloop()